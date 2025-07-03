from flask import Flask, render_template, request, redirect, url_for, flash, send_file, jsonify, session
from werkzeug.utils import secure_filename
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.exceptions import RequestEntityTooLarge
from datetime import datetime, timedelta
import os
import uuid
import zipfile
import json
import qrcode
import base64
from io import BytesIO
import tempfile
import shutil
from PIL import Image
# import fitz  # PyMuPDF - temporariamente desabilitado
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io
import yt_dlp
import threading
from config import config

# Configurar a aplicação baseada no ambiente
env = os.environ.get('FLASK_ENV', 'development')
app = Flask(__name__)

# Só configurar automaticamente se não estiver em produção
if env != 'production':
    app.config.from_object(config[env])
# Configurar diretório de uploads
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

# Modelos do banco de dados
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(120), nullable=False)
    two_factor_secret = db.Column(db.String(32), nullable=True, default=None)
    two_factor_enabled = db.Column(db.Boolean, default=False, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    files = db.relationship('File', backref='owner', lazy=True)
    shared_files = db.relationship('FileShare', backref='shared_by', lazy=True, foreign_keys='FileShare.shared_by_id')
    received_shares = db.relationship('FileShare', backref='shared_with', lazy=True, foreign_keys='FileShare.shared_with_id')
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        # Garantir que as colunas 2FA tenham valores padrão
        if not hasattr(self, 'two_factor_secret'):
            self.two_factor_secret = None
        if not hasattr(self, 'two_factor_enabled'):
            self.two_factor_enabled = False

class File(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255), nullable=False)
    original_filename = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(500), nullable=False)
    file_size = db.Column(db.Integer, nullable=False)
    file_type = db.Column(db.String(50), nullable=False)
    is_folder = db.Column(db.Boolean, default=False)
    parent_folder_id = db.Column(db.Integer, db.ForeignKey('file.id'), nullable=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Novos campos para tags e favoritos
    is_favorite = db.Column(db.Boolean, default=False)
    tags = db.Column(db.String(500), nullable=True)  # Tags separadas por vírgula
    last_accessed = db.Column(db.DateTime, nullable=True)  # Para arquivos recentes
    
    # Campos para lixeira
    is_deleted = db.Column(db.Boolean, default=False)
    deleted_at = db.Column(db.DateTime, nullable=True)
    
    # Relacionamento para pastas
    children = db.relationship('File', backref=db.backref('parent', remote_side=[id]))
    shares = db.relationship('FileShare', backref='file', lazy=True)
    
    # Campo para pasta compartilhada da organização
    shared_folder_id = db.Column(db.Integer, db.ForeignKey('shared_folder.id'), nullable=True)

class FileShare(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    file_id = db.Column(db.Integer, db.ForeignKey('file.id'), nullable=False)
    shared_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    shared_with_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    permission = db.Column(db.String(20), default='view')  # 'view', 'edit', 'admin'
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    expires_at = db.Column(db.DateTime, nullable=True)
    is_active = db.Column(db.Boolean, default=True)

# Modelos para Organizações e Equipes
class Organization(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=True)
    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relacionamentos
    owner = db.relationship('User', backref='owned_organizations', foreign_keys=[owner_id])
    teams = db.relationship('Team', backref='organization', lazy=True, cascade='all, delete-orphan')
    projects = db.relationship('Project', backref='organization', lazy=True, cascade='all, delete-orphan')
    shared_folders = db.relationship('SharedFolder', backref='organization', lazy=True, cascade='all, delete-orphan')

class Team(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=True)
    organization_id = db.Column(db.Integer, db.ForeignKey('organization.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relacionamentos
    members = db.relationship('TeamMember', backref='team', lazy=True, cascade='all, delete-orphan')

class TeamMember(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    team_id = db.Column(db.Integer, db.ForeignKey('team.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    role = db.Column(db.String(50), default='member')  # 'member', 'admin', 'owner'
    joined_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relacionamentos
    user = db.relationship('User', backref='team_memberships')

class Project(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=True)
    organization_id = db.Column(db.Integer, db.ForeignKey('organization.id'), nullable=False)
    created_by = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relacionamentos
    creator = db.relationship('User', backref='created_projects', foreign_keys=[created_by])

class SharedFolder(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=True)
    organization_id = db.Column(db.Integer, db.ForeignKey('organization.id'), nullable=False)
    created_by = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relacionamentos
    creator = db.relationship('User', backref='created_shared_folders', foreign_keys=[created_by])
    files = db.relationship('File', backref='shared_folder', lazy=True, foreign_keys='File.shared_folder_id')

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(e):
    flash('Arquivo muito grande! O tamanho máximo permitido é 500MB.', 'error')
    return redirect(request.referrer or url_for('dashboard'))

def get_file_icon(filename):
    """Retorna o ícone baseado na extensão do arquivo"""
    ext = os.path.splitext(filename)[1].lower()
    icons = {
        '.pdf': '📄',
        '.doc': '📝', '.docx': '📝',
        '.xls': '📊', '.xlsx': '📊',
        '.ppt': '📈', '.pptx': '📈',
        '.txt': '📄',
        '.jpg': '🖼️', '.jpeg': '🖼️', '.png': '🖼️', '.gif': '🖼️',
        '.mp3': '🎵', '.wav': '🎵',
        '.mp4': '🎬', '.avi': '🎬', '.mov': '🎬',
        '.zip': '📦', '.rar': '📦', '.7z': '📦',
        '.py': '🐍', '.js': '📜', '.html': '🌐', '.css': '🎨'
    }
    return icons.get(ext, '📄')

def format_file_size(size_bytes):
    """Formata o tamanho do arquivo"""
    if size_bytes == 0:
        return "0 B"
    size_names = ["B", "KB", "MB", "GB", "TB"]
    import math
    i = int(math.floor(math.log(size_bytes, 1024)))
    p = math.pow(1024, i)
    s = round(size_bytes / p, 2)
    return f"{s} {size_names[i]}"

def can_preview_file(file_path):
    """Verifica se o arquivo pode ser visualizado no navegador"""
    ext = os.path.splitext(file_path)[1].lower()
    previewable_extensions = {
        '.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml',
        '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg',
        '.pdf'
    }
    return ext in previewable_extensions

def get_file_content(file_path):
    """Lê o conteúdo de um arquivo de texto"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except:
            return "Não foi possível ler o conteúdo do arquivo (arquivo binário ou codificação não suportada)"
    except Exception as e:
        return f"Erro ao ler arquivo: {str(e)}"

def get_language_class(filename):
    """Retorna a classe de linguagem para syntax highlighting"""
    ext = os.path.splitext(filename)[1].lower()
    language_map = {
        '.py': 'python',
        '.js': 'javascript',
        '.html': 'markup',
        '.css': 'css',
        '.json': 'json',
        '.xml': 'markup',
        '.md': 'markdown',
        '.txt': 'text'
    }
    return language_map.get(ext, 'text')

def check_file_permission(file_id, user_id, required_permission='view'):
    """
    Verifica se um usuário tem permissão para um arquivo
    Retorna: (tem_permissao, permissao_atual, eh_dono)
    """
    file = File.query.get(file_id)
    if not file:
        return False, None, False
    
    # Se é o dono do arquivo
    if file.user_id == user_id:
        return True, 'admin', True
    
    # Verificar compartilhamento
    share = FileShare.query.filter_by(
        file_id=file_id,
        shared_with_id=user_id,
        is_active=True
    ).filter(
        (FileShare.expires_at.is_(None)) | (FileShare.expires_at > datetime.utcnow())
    ).first()
    
    if not share:
        return False, None, False
    
    # Mapear permissões
    permission_levels = {
        'view': 1,
        'edit': 2,
        'admin': 3
    }
    
    required_level = permission_levels.get(required_permission, 1)
    user_level = permission_levels.get(share.permission, 1)
    
    return user_level >= required_level, share.permission, False

def get_user_file_permission(file_id, user_id):
    """
    Retorna a permissão atual do usuário para um arquivo
    """
    file = File.query.get(file_id)
    if not file:
        return None
    
    # Se é o dono
    if file.user_id == user_id:
        return 'admin'
    
    # Verificar compartilhamento
    share = FileShare.query.filter_by(
        file_id=file_id,
        shared_with_id=user_id,
        is_active=True
    ).filter(
        (FileShare.expires_at.is_(None)) | (FileShare.expires_at > datetime.utcnow())
    ).first()
    
    return share.permission if share else None

# Context processors para disponibilizar funções nos templates
@app.context_processor
def utility_processor():
    return {
        'get_file_icon': get_file_icon,
        'format_file_size': format_file_size,
        'can_preview_file': can_preview_file,
        'get_language_class': get_language_class,
        'check_file_permission': check_file_permission,
        'get_user_file_permission': get_user_file_permission
    }

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return render_template('index.html')

@app.route('/init_db')
def init_db():
    """Inicializa o banco de dados"""
    try:
        with app.app_context():
            db.create_all()
        return 'Banco de dados inicializado com sucesso!'
    except Exception as e:
        return f'Erro ao inicializar banco: {str(e)}'

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        try:
            # Criar tabelas se não existirem
            with app.app_context():
                db.create_all()
            
            username = request.form['username']
            email = request.form['email']
            password = request.form['password']
            
            if User.query.filter_by(username=username).first():
                flash('Nome de usuário já existe!', 'error')
                return redirect(url_for('register'))
            
            if User.query.filter_by(email=email).first():
                flash('Email já cadastrado!', 'error')
                return redirect(url_for('register'))
            
            user = User(
                username=username,
                email=email,
                password_hash=generate_password_hash(password)
            )
            db.session.add(user)
            db.session.commit()
            
            flash('Conta criada com sucesso!', 'success')
            return redirect(url_for('login'))
        except Exception as e:
            print(f"Erro no registro: {e}")
            import traceback
            traceback.print_exc()
            flash('Erro interno. Tente novamente.', 'error')
            return redirect(url_for('register'))
    
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        try:
            user = User.query.filter_by(username=username).first()
            if user and check_password_hash(user.password_hash, password):
                # Verificar se as colunas 2FA existem e inicializar se necessário
                if not hasattr(user, 'two_factor_secret'):
                    user.two_factor_secret = None
                if not hasattr(user, 'two_factor_enabled'):
                    user.two_factor_enabled = False
                login_user(user)
                return redirect(url_for('dashboard'))
            else:
                flash('Usuário ou senha incorretos!', 'error')
        except Exception as e:
            print(f"Erro no login: {e}")
            flash('Erro interno. Tente novamente.', 'error')
    
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))

@app.route('/dashboard')
@login_required
def dashboard():
    try:
        folder_id = request.args.get('folder', type=int)
        
        if folder_id:
            current_folder = File.query.filter_by(id=folder_id, user_id=current_user.id).first()
            if not current_folder or not current_folder.is_folder:
                flash('Pasta não encontrada!', 'error')
                return redirect(url_for('dashboard'))
        else:
            current_folder = None
        
        # Buscar arquivos e pastas do usuário (não excluídos)
        query = File.query.filter_by(
            user_id=current_user.id,
            parent_folder_id=folder_id if folder_id else None,
            is_deleted=False
        ).order_by(File.is_folder.desc(), File.filename.asc())
        files = query.all()
        
        # Buscar arquivos compartilhados com o usuário (apenas na pasta raiz)
        shared_files = []
        if not folder_id:
            try:
                # Consulta melhorada para arquivos compartilhados
                from datetime import datetime
                shared_shares = FileShare.query.filter_by(
                    shared_with_id=current_user.id,
                    is_active=True
                ).filter(
                    (FileShare.expires_at.is_(None)) | (FileShare.expires_at > datetime.utcnow())
                ).all()
                for share in shared_shares:
                    try:
                        file = File.query.get(share.file_id)
                        if file and not file.parent_folder_id:  # Apenas arquivos na raiz
                            shared_files.append({
                                'file': file,
                                'share': share,
                                'is_shared': True
                            })
                    except Exception as e:
                        print(f"Erro ao processar compartilhamento {share.id}: {e}")
                        continue
            except Exception as e:
                print(f"Erro ao buscar arquivos compartilhados: {e}")
                shared_files = []
        
        # Calcular breadcrumb
        breadcrumb = []
        if current_folder:
            try:
                parent = current_folder.parent
                while parent:
                    breadcrumb.insert(0, parent)
                    parent = parent.parent
            except Exception as e:
                print(f"Erro ao calcular breadcrumb: {e}")
                breadcrumb = []
        
        return render_template('dashboard.html', 
                             files=files, 
                             shared_files=shared_files,
                             current_folder=current_folder,
                             breadcrumb=breadcrumb)
    except Exception as e:
        print(f"Erro no dashboard: {e}")
        import traceback
        traceback.print_exc()
        flash('Erro interno. Tente novamente.', 'error')
        return redirect(url_for('index'))

@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    if 'file' not in request.files:
        flash('Nenhum arquivo selecionado!', 'error')
        return redirect(request.referrer)
    
    files = request.files.getlist('file')
    folder_id = request.form.get('folder_id', type=int)
    
    # Criar pasta do usuário se não existir
    user_upload_folder = os.path.join(app.config['UPLOAD_FOLDER'], f'user_{current_user.id}')
    os.makedirs(user_upload_folder, exist_ok=True)
    
    for file in files:
        if file.filename == '':
            continue
        
        if file:
            # Verificar se a extensão é permitida
            if not allowed_file(file.filename):
                flash(f'Extensão não permitida para {file.filename}', 'error')
                continue
            
            # Gerar nome único para o arquivo
            filename = safe_filename(file.filename)
            unique_filename = f"{uuid.uuid4()}_{filename}"
            file_path = os.path.join(user_upload_folder, unique_filename)
            
            # Salvar arquivo
            file.save(file_path)
            
            # Salvar no banco de dados
            db_file = File(
                filename=unique_filename,
                original_filename=filename,
                file_path=file_path,
                file_size=os.path.getsize(file_path),
                file_type=os.path.splitext(filename)[1].lower(),
                user_id=current_user.id,
                parent_folder_id=folder_id
            )
            db.session.add(db_file)
    
    db.session.commit()
    flash('Arquivo(s) enviado(s) com sucesso!', 'success')
    return redirect(request.referrer)

@app.route('/upload_multiple', methods=['POST'])
@login_required
def upload_multiple_files():
    """Upload de múltiplos arquivos selecionados individualmente"""
    if 'files' not in request.files:
        flash('Nenhum arquivo selecionado!', 'error')
        return redirect(request.referrer)
    
    files = request.files.getlist('files')
    folder_id = request.form.get('folder_id', type=int)
    uploaded_count = 0
    
    # Criar pasta do usuário se não existir
    user_upload_folder = os.path.join(app.config['UPLOAD_FOLDER'], f'user_{current_user.id}')
    os.makedirs(user_upload_folder, exist_ok=True)
    
    for file in files:
        if file.filename == '':
            continue
        
        if file:
            try:
                # Verificar se a extensão é permitida
                if not allowed_file(file.filename):
                    flash(f'Extensão não permitida para {file.filename}', 'error')
                    continue
                
                # Gerar nome único para o arquivo
                filename = safe_filename(file.filename)
                unique_filename = f"{uuid.uuid4()}_{filename}"
                file_path = os.path.join(user_upload_folder, unique_filename)
                
                # Salvar arquivo
                file.save(file_path)
                
                # Salvar no banco de dados
                db_file = File(
                    filename=unique_filename,
                    original_filename=filename,
                    file_path=file_path,
                    file_size=os.path.getsize(file_path),
                    file_type=os.path.splitext(filename)[1].lower(),
                    user_id=current_user.id,
                    parent_folder_id=folder_id
                )
                db.session.add(db_file)
                uploaded_count += 1
                
            except Exception as e:
                flash(f'Erro ao enviar {file.filename}: {str(e)}', 'error')
                continue
    
    db.session.commit()
    flash(f'{uploaded_count} arquivo(s) enviado(s) com sucesso!', 'success')
    return redirect(request.referrer)

@app.route('/upload_folder', methods=['POST'])
@login_required
def upload_folder():
    """Upload de pasta completa mantendo a estrutura de diretórios"""
    if 'folder_files' not in request.files:
        flash('Nenhuma pasta selecionada!', 'error')
        return redirect(request.referrer)
    
    files = request.files.getlist('folder_files')
    parent_folder_id = request.form.get('folder_id', type=int)
    
    if not files:
        flash('Nenhum arquivo encontrado na pasta!', 'error')
        return redirect(request.referrer)
    
    # Criar pasta do usuário se não existir
    user_upload_folder = os.path.join(app.config['UPLOAD_FOLDER'], f'user_{current_user.id}')
    os.makedirs(user_upload_folder, exist_ok=True)
    
    # Organizar arquivos por estrutura de pastas
    folder_structure = {}
    for file in files:
        if file.filename == '':
            continue
        
        # Extrair caminho relativo da pasta
        path_parts = file.filename.split('/')
        if len(path_parts) > 1:
            folder_name = path_parts[0]
            if folder_name not in folder_structure:
                folder_structure[folder_name] = []
            folder_structure[folder_name].append({
                'file': file,
                'relative_path': '/'.join(path_parts[1:])
            })
    
    uploaded_count = 0
    created_folders = {}
    
    for folder_name, folder_files in folder_structure.items():
        # Criar pasta principal se não existir
        if folder_name not in created_folders:
            existing_folder = File.query.filter_by(
                filename=folder_name,
                user_id=current_user.id,
                parent_folder_id=parent_folder_id,
                is_folder=True
            ).first()
            
            if not existing_folder:
                folder = File(
                    filename=folder_name,
                    original_filename=folder_name,
                    file_path='',
                    file_size=0,
                    file_type='folder',
                    is_folder=True,
                    user_id=current_user.id,
                    parent_folder_id=parent_folder_id
                )
                db.session.add(folder)
                db.session.flush()  # Para obter o ID da pasta
                created_folders[folder_name] = folder.id
            else:
                created_folders[folder_name] = existing_folder.id
        
        # Upload dos arquivos da pasta
        for file_info in folder_files:
            file = file_info['file']
            relative_path = file_info['relative_path']
            
            try:
                # Verificar se há subpastas no caminho relativo
                path_parts = relative_path.split('/')
                current_parent_id = created_folders[folder_name]
                
                # Criar subpastas se necessário
                for i in range(len(path_parts) - 1):
                    subfolder_name = path_parts[i]
                    
                    # Verificar se a subpasta já existe
                    existing_subfolder = File.query.filter_by(
                        filename=subfolder_name,
                        user_id=current_user.id,
                        parent_folder_id=current_parent_id,
                        is_folder=True
                    ).first()
                    
                    if not existing_subfolder:
                        subfolder = File(
                            filename=subfolder_name,
                            original_filename=subfolder_name,
                            file_path='',
                            file_size=0,
                            file_type='folder',
                            is_folder=True,
                            user_id=current_user.id,
                            parent_folder_id=current_parent_id
                        )
                        db.session.add(subfolder)
                        db.session.flush()
                        current_parent_id = subfolder.id
                    else:
                        current_parent_id = existing_subfolder.id
                
                # Upload do arquivo
                filename = secure_filename(path_parts[-1])
                unique_filename = f"{uuid.uuid4()}_{filename}"
                file_path = os.path.join(user_upload_folder, unique_filename)
                
                file.save(file_path)
                
                db_file = File(
                    filename=unique_filename,
                    original_filename=filename,
                    file_path=file_path,
                    file_size=os.path.getsize(file_path),
                    file_type=os.path.splitext(filename)[1].lower(),
                    user_id=current_user.id,
                    parent_folder_id=current_parent_id
                )
                db.session.add(db_file)
                uploaded_count += 1
                
            except Exception as e:
                flash(f'Erro ao enviar {file.filename}: {str(e)}', 'error')
                continue
    
    db.session.commit()
    flash(f'Pasta enviada com sucesso! {uploaded_count} arquivo(s) processado(s).', 'success')
    return redirect(request.referrer)

@app.route('/create_folder', methods=['POST'])
@login_required
def create_folder():
    folder_name = request.form.get('folder_name')
    parent_folder_id = request.form.get('parent_folder_id', type=int)
    
    if not folder_name:
        flash('Nome da pasta é obrigatório!', 'error')
        return redirect(request.referrer)
    
    # Verificar se já existe uma pasta com esse nome
    existing = File.query.filter_by(
        filename=folder_name,
        user_id=current_user.id,
        parent_folder_id=parent_folder_id,
        is_folder=True
    ).first()
    
    if existing:
        flash('Já existe uma pasta com esse nome!', 'error')
        return redirect(request.referrer)
    
    # Criar pasta no banco de dados
    folder = File(
        filename=folder_name,
        original_filename=folder_name,
        file_path='',  # Pastas não têm arquivo físico
        file_size=0,
        file_type='folder',
        is_folder=True,
        user_id=current_user.id,
        parent_folder_id=parent_folder_id
    )
    
    db.session.add(folder)
    db.session.commit()
    
    flash('Pasta criada com sucesso!', 'success')
    return redirect(request.referrer)

@app.route('/download/<int:file_id>')
@login_required
def download_file(file_id):
    # Verificar permissão
    has_permission, permission, is_owner = check_file_permission(file_id, current_user.id, 'view')
    if not has_permission:
        flash('Você não tem permissão para baixar este arquivo!', 'error')
        return redirect(url_for('dashboard'))
    
    file = File.query.get_or_404(file_id)
    
    if file.is_folder:
        flash('Não é possível baixar uma pasta!', 'error')
        return redirect(url_for('dashboard'))
    
    try:
        return send_file(file.file_path, 
                        as_attachment=True, 
                        download_name=file.original_filename)
    except Exception as e:
        flash(f'Erro ao baixar arquivo: {str(e)}', 'error')
        return redirect(url_for('dashboard'))

@app.route('/preview/<int:file_id>')
@login_required
def preview_file(file_id):
    # Verificar permissão
    has_permission, permission, is_owner = check_file_permission(file_id, current_user.id, 'view')
    if not has_permission:
        flash('Você não tem permissão para visualizar este arquivo!', 'error')
        return redirect(url_for('dashboard'))
    
    file = File.query.get_or_404(file_id)
    
    if file.is_folder:
        flash('Não é possível visualizar uma pasta!', 'error')
        return redirect(url_for('dashboard'))
    
    if not can_preview_file(file.file_path):
        flash('Este tipo de arquivo não pode ser visualizado!', 'error')
        return redirect(url_for('dashboard'))
    
    ext = os.path.splitext(file.file_path)[1].lower()
    
    # Para imagens, retornar diretamente
    if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg']:
        return send_file(file.file_path)
    
    # Para PDFs, retornar diretamente
    elif ext == '.pdf':
        return send_file(file.file_path, mimetype='application/pdf', as_attachment=False)
    
    # Para arquivos de texto, mostrar em template
    elif ext in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml']:
        content = get_file_content(file.file_path)
        return render_template('preview.html', file=file, content=content)
    
    else:
        flash('Tipo de arquivo não suportado para visualização!', 'error')
        return redirect(url_for('dashboard'))

@app.route('/view/<int:file_id>')
@login_required
def view_file(file_id):
    # Verificar permissão
    has_permission, permission, is_owner = check_file_permission(file_id, current_user.id, 'view')
    if not has_permission:
        flash('Você não tem permissão para visualizar este arquivo!', 'error')
        return redirect(url_for('dashboard'))
    
    file = File.query.get_or_404(file_id)
    
    # Atualizar último acesso (apenas para arquivos, não pastas)
    if not file.is_folder:
        file.last_accessed = datetime.utcnow()
        db.session.commit()
    
    # Redirecionar arquivos de mídia para o player
    media_extensions = ['.mp4', '.avi', '.mov', '.mkv', '.webm', '.mp3', '.wav', '.ogg', '.flac', '.aac']
    if file.file_type in media_extensions:
        return redirect(url_for('view_media', file_id=file_id))
    
    # Verificar se é um arquivo editável
    editable_extensions = ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml']
    if file.file_type in editable_extensions:
        return redirect(url_for('edit_file', file_id=file_id))
    
    # Verificar se é uma imagem - renderizar template de carrossel
    image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']
    if file.file_type in image_extensions:
        # Salvar referrer na sessão para o carrossel
        session['view_referrer'] = request.referrer or url_for('dashboard')
        return render_template('view_image.html', file=file)
    
    # Verificar se é um PDF
    if file.file_type == '.pdf':
        return redirect(url_for('preview_file', file_id=file_id))
    
    # Para outros tipos de arquivo, mostrar conteúdo como texto
    try:
        content = get_file_content(file.file_path)
        return render_template('view_text.html', file=file, content=content)
    except Exception as e:
        flash(f'Erro ao ler arquivo: {str(e)}', 'error')
        return redirect(url_for('dashboard'))

@app.route('/delete/<int:file_id>', methods=['POST'])
@login_required
def delete_file(file_id):
    # Verificar permissão de admin
    has_permission, permission, is_owner = check_file_permission(file_id, current_user.id, 'admin')
    if not has_permission:
        flash('Você não tem permissão para excluir este arquivo!', 'error')
        return redirect(url_for('dashboard'))
    
    file = File.query.get_or_404(file_id)
    
    # Marcar como excluído (mover para lixeira)
    file.is_deleted = True
    file.deleted_at = datetime.utcnow()
    
    # Desativar compartilhamentos relacionados ao arquivo
    FileShare.query.filter_by(file_id=file_id).update({'is_active': False})
    
    db.session.commit()
    
    flash('Arquivo movido para a lixeira!', 'success')
    return redirect(request.referrer)

@app.route('/edit/<int:file_id>', methods=['GET', 'POST'])
@login_required
def edit_file(file_id):
    # Verificar permissão de edição
    has_permission, permission, is_owner = check_file_permission(file_id, current_user.id, 'edit')
    if not has_permission:
        flash('Você não tem permissão para editar este arquivo!', 'error')
        return redirect(url_for('dashboard'))
    
    file = File.query.get_or_404(file_id)
    
    if file.is_folder:
        flash('Não é possível editar uma pasta!', 'error')
        return redirect(url_for('dashboard'))
    
    # Verificar se é um arquivo de texto editável
    ext = os.path.splitext(file.file_path)[1].lower()
    editable_extensions = ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml']
    
    if ext not in editable_extensions:
        flash('Este tipo de arquivo não pode ser editado!', 'error')
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        new_content = request.form.get('content')
        if new_content is not None:
            try:
                # Salvar novo conteúdo
                with open(file.file_path, 'w', encoding='utf-8') as f:
                    f.write(new_content)
                
                # Atualizar tamanho do arquivo
                file.file_size = os.path.getsize(file.file_path)
                file.updated_at = datetime.utcnow()
                db.session.commit()
                
                flash('Arquivo salvo com sucesso!', 'success')
                return redirect(url_for('view_file', file_id=file_id))
            except Exception as e:
                flash(f'Erro ao salvar arquivo: {str(e)}', 'error')
    
    # Carregar conteúdo atual
    try:
        with open(file.file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        flash(f'Erro ao ler arquivo: {str(e)}', 'error')
        return redirect(url_for('dashboard'))
    
    return render_template('edit_file.html', file=file, content=content)

@app.route('/api/storage_info')
@login_required
def storage_info():
    """Retorna informações de armazenamento do usuário"""
    try:
        print(f"API Debug - Usuário: {current_user.id}")
        
        files = File.query.filter_by(user_id=current_user.id).all()
        print(f"API Debug - Total de arquivos encontrados: {len(files)}")
        
        # Filtrar apenas arquivos (não pastas)
        file_files = [f for f in files if not f.is_folder]
        total_size = sum(file.file_size for file in file_files)
        file_count = len(file_files)
        
        print(f"API Debug - Arquivos (não pastas): {file_count}")
        print(f"API Debug - Tamanho total: {total_size} bytes")
        
        formatted_size = format_file_size(total_size)
        print(f"API Debug - Tamanho formatado: {formatted_size}")
        
        result = {
            'total_size': formatted_size,
            'file_count': file_count,
            'debug': {
                'user_id': current_user.id,
                'total_files_in_db': len(files),
                'files_only': file_count,
                'total_size_bytes': total_size
            }
        }
        
        print(f"API Debug - Resultado: {result}")
        return jsonify(result)
        
    except Exception as e:
        print(f"API Debug - Erro: {str(e)}")
        return jsonify({
            'error': str(e),
            'total_size': '0 B',
            'file_count': 0
        }), 500

@app.route('/settings')
@login_required
def settings():
    """Página de configurações do usuário"""
    return render_template('settings.html')

@app.route('/settings/change_password', methods=['POST'])
@login_required
def change_password():
    """Alterar senha do usuário"""
    current_password = request.form.get('current_password')
    new_password = request.form.get('new_password')
    confirm_password = request.form.get('confirm_password')
    
    if not check_password_hash(current_user.password_hash, current_password):
        flash('Senha atual incorreta!', 'error')
        return redirect(url_for('settings'))
    
    if new_password != confirm_password:
        flash('As senhas não coincidem!', 'error')
        return redirect(url_for('settings'))
    
    if len(new_password) < 6:
        flash('A nova senha deve ter pelo menos 6 caracteres!', 'error')
        return redirect(url_for('settings'))
    
    current_user.password_hash = generate_password_hash(new_password)
    db.session.commit()
    
    flash('Senha alterada com sucesso!', 'success')
    return redirect(url_for('settings'))

@app.route('/settings/update_profile', methods=['POST'])
@login_required
def update_profile():
    """Atualizar perfil do usuário"""
    username = request.form.get('username')
    email = request.form.get('email')
    
    # Verificar se o username já existe (exceto o atual)
    existing_user = User.query.filter_by(username=username).first()
    if existing_user and existing_user.id != current_user.id:
        flash('Nome de usuário já existe!', 'error')
        return redirect(url_for('settings'))
    
    # Verificar se o email já existe (exceto o atual)
    existing_email = User.query.filter_by(email=email).first()
    if existing_email and existing_email.id != current_user.id:
        flash('Email já cadastrado!', 'error')
        return redirect(url_for('settings'))
    
    current_user.username = username
    current_user.email = email
    db.session.commit()
    
    flash('Perfil atualizado com sucesso!', 'success')
    return redirect(url_for('settings'))

@app.route('/settings/export_data')
@login_required
def export_data():
    """Exportar dados do usuário"""
    # Coletar dados do usuário
    user_data = {
        'username': current_user.username,
        'email': current_user.email,
        'created_at': current_user.created_at.isoformat(),
        'files': []
    }
    
    # Coletar informações dos arquivos
    files = File.query.filter_by(user_id=current_user.id).all()
    for file in files:
        file_info = {
            'filename': file.original_filename,
            'file_type': file.file_type,
            'file_size': file.file_size,
            'is_folder': file.is_folder,
            'created_at': file.created_at.isoformat(),
            'updated_at': file.updated_at.isoformat()
        }
        user_data['files'].append(file_info)
    
    # Criar resposta JSON
    response = jsonify(user_data)
    response.headers['Content-Disposition'] = f'attachment; filename=user_data_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json'
    return response

@app.route('/settings/delete_account', methods=['POST'])
@login_required
def delete_account():
    """Deletar conta do usuário"""
    password = request.form.get('password')
    
    if not check_password_hash(current_user.password_hash, password):
        flash('Senha incorreta!', 'error')
        return redirect(url_for('settings'))
    
    # Deletar todos os arquivos do usuário
    files = File.query.filter_by(user_id=current_user.id).all()
    for file in files:
        try:
            if os.path.exists(file.file_path):
                os.remove(file.file_path)
        except:
            pass
    
    # Deletar usuário do banco
    db.session.delete(current_user)
    db.session.commit()
    
    logout_user()
    flash('Conta deletada com sucesso!', 'success')
    return redirect(url_for('index'))

@app.route('/settings/setup_2fa')
@login_required
def setup_2fa():
    """Configurar autenticação de dois fatores"""
    if current_user.two_factor_enabled:
        flash('2FA já está habilitado!', 'info')
        return redirect(url_for('settings'))
    
    # Gerar chave secreta
    secret = secrets.token_hex(16)
    current_user.two_factor_secret = secret
    db.session.commit()
    
    # Gerar QR Code
    qr = qrcode.QRCode(version=1, box_size=10, border=5)
    qr.add_data(f'otpauth://totp/CloudStorage:{current_user.username}?secret={secret}&issuer=CloudStorage')
    qr.make(fit=True)
    
    img = qr.make_image(fill_color="black", back_color="white")
    
    # Converter para base64
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    qr_code_b64 = base64.b64encode(buffer.getvalue()).decode()
    
    return render_template('setup_2fa.html', qr_code=f'data:image/png;base64,{qr_code_b64}', secret=secret)

@app.route('/settings/verify_2fa', methods=['POST'])
@login_required
def verify_2fa():
    """Verificar código 2FA e habilitar"""
    code = request.form.get('code')
    
    if not code or len(code) != 6:
        flash('Código inválido!', 'error')
        return redirect(url_for('setup_2fa'))
    
    # Aqui você implementaria a verificação real do TOTP
    # Por simplicidade, vamos aceitar qualquer código de 6 dígitos
    if code.isdigit() and len(code) == 6:
        current_user.two_factor_enabled = True
        db.session.commit()
        flash('Autenticação de dois fatores habilitada com sucesso!', 'success')
        return redirect(url_for('settings'))
    else:
        flash('Código inválido!', 'error')
        return redirect(url_for('setup_2fa'))

@app.route('/settings/disable_2fa', methods=['POST'])
@login_required
def disable_2fa():
    """Desabilitar autenticação de dois fatores"""
    password = request.form.get('password')
    
    if not check_password_hash(current_user.password_hash, password):
        flash('Senha incorreta!', 'error')
        return redirect(url_for('settings'))
    
    current_user.two_factor_enabled = False
    current_user.two_factor_secret = None
    db.session.commit()
    
    flash('Autenticação de dois fatores desabilitada!', 'success')
    return redirect(url_for('settings'))

@app.route('/reports')
@login_required
def reports():
    """Página de relatórios avançados"""
    # Estatísticas gerais
    files = File.query.filter_by(user_id=current_user.id).all()
    total_files = len([f for f in files if not f.is_folder])
    total_folders = len([f for f in files if f.is_folder])
    total_size = sum(f.file_size for f in files if not f.is_folder)
    
    # Estatísticas por tipo de arquivo
    file_types = {}
    for file in files:
        if not file.is_folder:
            ext = file.file_type.lower()
            if ext not in file_types:
                file_types[ext] = {'count': 0, 'size': 0}
            file_types[ext]['count'] += 1
            file_types[ext]['size'] += file.file_size
    
    # Top 10 tipos de arquivo
    top_file_types = sorted(file_types.items(), key=lambda x: x[1]['size'], reverse=True)[:10]
    
    # Estatísticas por mês
    from collections import defaultdict
    monthly_stats = defaultdict(lambda: {'count': 0, 'size': 0})
    
    for file in files:
        if not file.is_folder:
            month = file.created_at.strftime('%Y-%m')
            monthly_stats[month]['count'] += 1
            monthly_stats[month]['size'] += file.file_size
    
    # Ordenar por mês
    monthly_data = sorted(monthly_stats.items())
    
    # Arquivos recentes
    recent_files = File.query.filter_by(user_id=current_user.id, is_folder=False)\
                             .order_by(File.created_at.desc()).limit(10).all()
    
    # Maiores arquivos
    largest_files = File.query.filter_by(user_id=current_user.id, is_folder=False)\
                              .order_by(File.file_size.desc()).limit(10).all()
    
    return render_template('reports.html',
                         total_files=total_files,
                         total_folders=total_folders,
                         total_size=total_size,
                         top_file_types=top_file_types,
                         monthly_data=monthly_data,
                         recent_files=recent_files,
                         largest_files=largest_files)

# Rotas de Compartilhamento
@app.route('/share/<int:file_id>', methods=['GET', 'POST'])
@login_required
def share_file(file_id):
    file = File.query.get_or_404(file_id)
    
    # Verificar se o usuário é dono do arquivo
    if file.user_id != current_user.id:
        flash('Você só pode compartilhar seus próprios arquivos!', 'error')
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        permission = request.form.get('permission', 'view')
        expires_in = request.form.get('expires_in', 'never')
        
        # Buscar usuário
        user = User.query.filter_by(username=username).first()
        if not user:
            flash('Usuário não encontrado!', 'error')
            return redirect(url_for('share_file', file_id=file_id))
        
        if user.id == current_user.id:
            flash('Você não pode compartilhar com você mesmo!', 'error')
            return redirect(url_for('share_file', file_id=file_id))
        
        # Verificar se já existe compartilhamento
        existing_share = FileShare.query.filter_by(
            file_id=file_id,
            shared_with_id=user.id,
            is_active=True
        ).first()
        
        if existing_share:
            flash(f'Arquivo já está compartilhado com {username}!', 'error')
            return redirect(url_for('share_file', file_id=file_id))
        
        # Calcular data de expiração
        expires_at = None
        if expires_in != 'never':
            from datetime import timedelta
            if expires_in == '1_day':
                expires_at = datetime.utcnow() + timedelta(days=1)
            elif expires_in == '1_week':
                expires_at = datetime.utcnow() + timedelta(weeks=1)
            elif expires_in == '1_month':
                expires_at = datetime.utcnow() + timedelta(days=30)
        
        # Criar compartilhamento
        share = FileShare(
            file_id=file_id,
            shared_by_id=current_user.id,
            shared_with_id=user.id,
            permission=permission,
            expires_at=expires_at
        )
        
        db.session.add(share)
        db.session.commit()
        
        flash(f'Arquivo compartilhado com {username} com sucesso!', 'success')
        return redirect(url_for('dashboard'))
    
    # Buscar compartilhamentos existentes
    shares = FileShare.query.filter_by(file_id=file_id, is_active=True).all()
    users = User.query.filter(User.id != current_user.id).all()
    
    return render_template('share_file.html', file=file, shares=shares, users=users)

@app.route('/shared_with_me')
@login_required
def shared_with_me():
    """Arquivos compartilhados com o usuário atual"""
    try:
        shares = FileShare.query.filter_by(
            shared_with_id=current_user.id,
            is_active=True
        ).all()
        
        return render_template('shared_with_me.html', shares=shares)
    except Exception as e:
        print(f"Erro em shared_with_me: {e}")
        flash('Erro interno. Tente novamente.', 'error')
        return redirect(url_for('dashboard'))

@app.route('/shared_by_me')
@login_required
def shared_by_me():
    """Página de arquivos que o usuário compartilhou"""
    shares = FileShare.query.filter_by(
        shared_by_id=current_user.id,
        is_active=True
    ).order_by(FileShare.created_at.desc()).all()
    
    return render_template('shared_by_me.html', shares=shares)

@app.route('/manage_shared_files')
@login_required
def manage_shared_files():
    """Página para gerenciar arquivos compartilhados"""
    # Arquivos que eu compartilhei
    my_shares = FileShare.query.filter_by(
        shared_by_id=current_user.id,
        is_active=True
    ).order_by(FileShare.created_at.desc()).all()
    
    # Arquivos compartilhados comigo
    received_shares = FileShare.query.filter_by(
        shared_with_id=current_user.id,
        is_active=True
    ).filter(
        (FileShare.expires_at.is_(None)) | (FileShare.expires_at > datetime.utcnow())
    ).order_by(FileShare.created_at.desc()).all()
    
    return render_template('manage_shared_files.html', 
                         my_shares=my_shares, 
                         received_shares=received_shares)

@app.route('/revoke_share/<int:share_id>', methods=['POST'])
@login_required
def revoke_share(share_id):
    share = FileShare.query.get_or_404(share_id)
    
    # Verificar se o usuário é quem compartilhou
    if share.shared_by_id != current_user.id:
        flash('Você só pode revogar seus próprios compartilhamentos!', 'error')
        return redirect(url_for('shared_by_me'))
    
    share.is_active = False
    db.session.commit()
    
    flash('Compartilhamento revogado com sucesso!', 'success')
    return redirect(url_for('shared_by_me'))

@app.route('/api/users')
@login_required
def get_users():
    """API para buscar usuários para autocomplete"""
    query = request.args.get('q', '')
    users = User.query.filter(
        User.username.like(f'%{query}%'),
        User.id != current_user.id
    ).limit(10).all()
    
    return jsonify([{'id': user.id, 'username': user.username} for user in users])

@app.route('/api/folder_images/<int:folder_id>')
@login_required
def get_folder_images(folder_id):
    """API para obter todas as imagens de uma pasta para o carrossel"""
    try:
        # Verificar se a pasta existe e pertence ao usuário
        if folder_id == 0:  # Pasta raiz
            folder = None
        else:
            folder = File.query.filter_by(id=folder_id, user_id=current_user.id, is_folder=True).first()
            if not folder:
                return jsonify({'error': 'Pasta não encontrada'}), 404
        
        # Buscar todas as imagens na pasta
        query = File.query.filter_by(
            user_id=current_user.id,
            parent_folder_id=folder_id if folder_id else None
        ).filter(
            File.is_folder == False,
            File.file_type.in_(['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg'])
        ).order_by(File.filename.asc())
        
        images = query.all()
        
        # Formatar dados das imagens
        image_list = []
        for img in images:
            image_list.append({
                'id': img.id,
                'filename': img.original_filename,
                'file_size': img.file_size,
                'created_at': img.created_at.isoformat(),
                'url': url_for('preview_file', file_id=img.id)
            })
        
        return jsonify({
            'images': image_list,
            'total': len(image_list)
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/folder_images')
@login_required
def get_current_folder_images():
    """API para obter imagens da pasta atual baseado no referrer"""
    try:
        referrer = request.args.get('referrer', '')
        print(f"API Debug - Referrer: {referrer}")
        
        # Extrair folder_id do referrer se existir
        folder_id = None
        if 'folder=' in referrer:
            import re
            match = re.search(r'folder=(\d+)', referrer)
            if match:
                folder_id = int(match.group(1))
                print(f"API Debug - Folder ID extraído: {folder_id}")
        
        print(f"API Debug - Buscando imagens na pasta: {folder_id}")
        
        # Buscar imagens na pasta
        query = File.query.filter_by(
            user_id=current_user.id,
            parent_folder_id=folder_id
        ).filter(
            File.is_folder == False,
            File.file_type.in_(['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg'])
        ).order_by(File.filename.asc())
        
        images = query.all()
        print(f"API Debug - Imagens encontradas: {len(images)}")
        
        # Formatar dados das imagens
        image_list = []
        for img in images:
            image_list.append({
                'id': img.id,
                'filename': img.original_filename,
                'file_size': img.file_size,
                'created_at': img.created_at.isoformat(),
                'url': url_for('preview_file', file_id=img.id)
            })
        
        return jsonify({
            'images': image_list,
            'total': len(image_list),
            'current_folder': folder_id
        })
        
    except Exception as e:
        print(f"API Debug - Erro: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/view_media/<int:file_id>')
@login_required
def view_media(file_id):
    """Visualizar arquivos de mídia com player integrado"""
    # Verificar permissão
    has_permission, permission, is_owner = check_file_permission(file_id, current_user.id, 'view')
    if not has_permission:
        flash('Você não tem permissão para visualizar este arquivo!', 'error')
        return redirect(url_for('dashboard'))
    
    file = File.query.get_or_404(file_id)
    
    # Verificar se é um arquivo de mídia
    media_extensions = ['.mp4', '.avi', '.mov', '.mkv', '.webm', '.mp3', '.wav', '.ogg', '.flac', '.aac']
    if file.file_type not in media_extensions:
        flash('Este arquivo não é um arquivo de mídia suportado!', 'error')
        return redirect(url_for('dashboard'))
    
    # Buscar outros arquivos de mídia na mesma pasta
    folder_media = []
    if file.parent_folder_id:
        folder_media = File.query.filter(
            File.parent_folder_id == file.parent_folder_id,
            File.file_type.in_(media_extensions),
            File.id != file.id
        ).limit(6).all()
    else:
        # Se estiver na raiz, buscar outros arquivos de mídia do usuário
        folder_media = File.query.filter(
            File.user_id == current_user.id,
            File.parent_folder_id.is_(None),
            File.file_type.in_(media_extensions),
            File.id != file.id
        ).limit(6).all()
    
    return render_template('view_media.html', file=file, folder_media=folder_media)

@app.route('/convert_files')
@login_required
def convert_files():
    """Página do conversor de formatos"""
    # Buscar pastas do usuário para organização
    user_folders = File.query.filter_by(
        user_id=current_user.id,
        is_folder=True
    ).order_by(File.filename.asc()).all()
    
    return render_template('convert_files.html', user_folders=user_folders)

@app.route('/api/convert_file', methods=['POST'])
@login_required
def convert_file():
    """API para conversão de arquivos"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'Nenhum arquivo enviado'}), 400
        
        file = request.files['file']
        target_format = request.form.get('target_format', '').lower()
        quality = request.form.get('quality', 'medium')
        folder_id = request.form.get('folder_id', '')
        
        if not file or file.filename == '':
            return jsonify({'error': 'Arquivo inválido'}), 400
        
        if not target_format:
            return jsonify({'error': 'Formato de destino não especificado'}), 400
        
        # Verificar pasta de destino
        destination_folder = None
        if folder_id:
            destination_folder = File.query.filter_by(
                id=folder_id, 
                user_id=current_user.id, 
                is_folder=True
            ).first()
            if not destination_folder:
                return jsonify({'error': 'Pasta de destino não encontrada'}), 404
        
        # Processar conversão
        result = process_file_conversion(file, target_format, quality, current_user.id, destination_folder)
        
        if result['success']:
            return jsonify(result)
        else:
            return jsonify({'error': result['error']}), 400
            
    except Exception as e:
        return jsonify({'error': f'Erro na conversão: {str(e)}'}), 500

def process_file_conversion(file, target_format, quality, user_id, destination_folder=None):
    """Processar conversão de arquivo"""
    try:
        import tempfile
        import os
        from PIL import Image
        import fitz  # PyMuPDF
        from docx import Document
        import io
        # NÃO importar pydub ou moviepy aqui!
        
        # Obter extensão original
        original_filename = file.filename
        file_ext = os.path.splitext(original_filename)[1].lower()
        
        # Verificar se a conversão é suportada
        if not get_supported_conversions(file_ext, target_format):
            return {'success': False, 'error': f'Conversão de {file_ext} para {target_format} não é suportada'}
        
        # Salvar arquivo temporário
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
            file.save(temp_file.name)
            temp_path = temp_file.name
        
        try:
            # Processar conversão baseada no tipo
            if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
                result = convert_image(temp_path, target_format, quality)
            elif file_ext in ['.mp4', '.avi', '.mov', '.mkv', '.webm']:
                result = convert_video(temp_path, target_format, quality)
            elif file_ext in ['.mp3', '.wav', '.ogg', '.flac', '.aac']:
                result = convert_audio(temp_path, target_format, quality)
            elif file_ext in ['.pdf', '.doc', '.docx', '.txt', '.rtf']:
                result = convert_document(temp_path, file_ext, target_format, quality)
            else:
                return {'success': False, 'error': f'Tipo de arquivo não suportado: {file_ext}'}
            
            if not result['success']:
                return result
            
            # Salvar arquivo convertido no sistema
            converted_filename = os.path.splitext(original_filename)[0] + '.' + target_format
            saved_file = save_converted_file(
                result['data'], 
                converted_filename, 
                user_id, 
                destination_folder
            )
            
            return {
                'success': True,
                'filename': converted_filename,
                'file_id': saved_file.id,
                'message': f'Arquivo convertido com sucesso para {target_format.upper()}'
            }
            
        finally:
            # Limpar arquivo temporário
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        return {'success': False, 'error': f'Erro na conversão: {str(e)}'}

def get_supported_conversions(source_ext, target_format):
    """Verificar se a conversão é suportada"""
    conversions = {
        '.jpg': ['png', 'webp', 'bmp', 'pdf'],
        '.jpeg': ['png', 'webp', 'bmp', 'pdf'],
        '.png': ['jpg', 'webp', 'bmp', 'pdf'],
        '.gif': ['jpg', 'png', 'webp', 'pdf'],
        '.bmp': ['jpg', 'png', 'webp', 'pdf'],
        '.webp': ['jpg', 'png', 'bmp', 'pdf'],
        '.mp4': ['avi', 'mov', 'mkv'],
        '.avi': ['mp4', 'mov', 'mkv'],
        '.mov': ['mp4', 'avi', 'mkv'],
        '.mkv': ['mp4', 'avi', 'mov'],
        '.webm': ['mp4', 'avi', 'mov'],
        '.mp3': ['wav', 'ogg', 'flac'],
        '.wav': ['mp3', 'ogg', 'flac'],
        '.ogg': ['mp3', 'wav', 'flac'],
        '.flac': ['mp3', 'wav', 'ogg'],
        '.aac': ['mp3', 'wav', 'ogg'],
        '.pdf': ['docx', 'doc', 'txt'],
        '.doc': ['pdf', 'txt'],
        '.docx': ['pdf', 'txt'],
        '.txt': ['pdf', 'docx'],
        '.rtf': ['pdf', 'docx', 'txt']
    }
    return target_format in conversions.get(source_ext, [])

def convert_image(file_path, target_format, quality):
    try:
        img = Image.open(file_path)
        quality_map = {'high': 95, 'medium': 85, 'low': 70}
        save_quality = quality_map.get(quality, 85)
        buffer = io.BytesIO()
        if target_format == 'jpg':
            if img.mode in ('RGBA', 'LA', 'P'):
                img = img.convert('RGB')
            img.save(buffer, format='JPEG', quality=save_quality)
        elif target_format == 'png':
            img.save(buffer, format='PNG', optimize=True)
        elif target_format == 'webp':
            img.save(buffer, format='WebP', quality=save_quality)
        elif target_format == 'bmp':
            img.save(buffer, format='BMP')
        elif target_format == 'pdf':
            img.save(buffer, format='PDF')
        else:
            return {'success': False, 'error': f'Formato de imagem não suportado: {target_format}'}
        buffer.seek(0)
        return {'success': True, 'data': buffer.getvalue()}
    except Exception as e:
        return {'success': False, 'error': f'Erro na conversão de imagem: {str(e)}'}

def convert_video(file_path, target_format, quality):
    try:
        # Conversão de vídeo temporariamente indisponível
        # Para habilitar, instale: pip install moviepy
        return {'success': False, 'error': f'Conversão de vídeo temporariamente indisponível. Formato solicitado: {target_format}. Para conversão de vídeo, instale moviepy: pip install moviepy'}
    except Exception as e:
        return {'success': False, 'error': f'Erro na conversão de vídeo: {str(e)}'}

def convert_audio(file_path, target_format, quality):
    try:
        # Conversão de áudio temporariamente indisponível
        # Para habilitar, instale: pip install pydub
        return {'success': False, 'error': f'Conversão de áudio temporariamente indisponível. Formato solicitado: {target_format}. Para conversão de áudio, instale pydub: pip install pydub'}
    except Exception as e:
        return {'success': False, 'error': f'Erro na conversão de áudio: {str(e)}'}

def convert_document(file_path, file_ext, target_format, quality):
    try:
        # PDF para DOCX
        if file_ext == '.pdf' and target_format in ['docx', 'doc']:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            docx_doc = Document()
            docx_doc.add_paragraph(text)
            buffer = io.BytesIO()
            docx_doc.save(buffer)
            buffer.seek(0)
            return {'success': True, 'data': buffer.getvalue()}
        # DOCX para PDF
        elif file_ext in ['.docx', '.doc'] and target_format == 'pdf':
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            doc = Document(file_path)
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            y = height - 40
            for para in doc.paragraphs:
                text = para.text
                c.drawString(40, y, text)
                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40
            c.save()
            buffer.seek(0)
            return {'success': True, 'data': buffer.getvalue()}
        # PDF para TXT
        elif file_ext == '.pdf' and target_format == 'txt':
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            buffer = io.BytesIO()
            buffer.write(text.encode('utf-8'))
            buffer.seek(0)
            return {'success': True, 'data': buffer.getvalue()}
        # TXT para PDF
        elif file_ext == '.txt' and target_format == 'pdf':
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            y = height - 40
            with open(file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    c.drawString(40, y, line.strip())
                    y -= 15
                    if y < 40:
                        c.showPage()
                        y = height - 40
            c.save()
            buffer.seek(0)
            return {'success': True, 'data': buffer.getvalue()}
        # TXT para DOCX
        elif file_ext == '.txt' and target_format == 'docx':
            doc = Document()
            with open(file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    doc.add_paragraph(line.strip())
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return {'success': True, 'data': buffer.getvalue()}
        # DOCX para TXT
        elif file_ext in ['.docx', '.doc'] and target_format == 'txt':
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            buffer = io.BytesIO()
            buffer.write(text.encode('utf-8'))
            buffer.seek(0)
            return {'success': True, 'data': buffer.getvalue()}
        else:
            return {'success': False, 'error': f'Conversão de documento não suportada: {file_ext} para {target_format}'}
    except Exception as e:
        return {'success': False, 'error': f'Erro na conversão de documento: {str(e)}'}

def save_converted_file(file_data, filename, user_id, destination_folder=None):
    try:
        import uuid
        import os
        unique_id = str(uuid.uuid4())
        file_extension = os.path.splitext(filename)[1]
        new_filename = f"{unique_id}_{filename}"
        if destination_folder:
            upload_folder = os.path.join('uploads', f'user_{user_id}', str(destination_folder.id))
        else:
            upload_folder = os.path.join('uploads', f'user_{user_id}')
        os.makedirs(upload_folder, exist_ok=True)
        file_path = os.path.join(upload_folder, new_filename)
        with open(file_path, 'wb') as f:
            f.write(file_data)
        file_size = len(file_data)
        new_file = File(
            filename=new_filename,
            original_filename=filename,
            file_path=file_path,
            file_size=file_size,
            file_type=file_extension,
            user_id=user_id,
            parent_folder_id=destination_folder.id if destination_folder else None
        )
        db.session.add(new_file)
        db.session.commit()
        return new_file
    except Exception as e:
        raise Exception(f'Erro ao salvar arquivo convertido: {str(e)}')

@app.route('/trash')
@login_required
def trash():
    """Página da lixeira - arquivos excluídos"""
    # Buscar arquivos excluídos do usuário
    deleted_files = File.query.filter_by(
        user_id=current_user.id,
        is_deleted=True
    ).order_by(File.deleted_at.desc()).all()
    
    return render_template('trash.html', files=deleted_files)

@app.route('/restore/<int:file_id>', methods=['POST'])
@login_required
def restore_file(file_id):
    """Restaurar arquivo da lixeira"""
    file = File.query.get_or_404(file_id)
    
    # Verificar se o arquivo pertence ao usuário
    if file.user_id != current_user.id:
        flash('Você não tem permissão para restaurar este arquivo!', 'error')
        return redirect(url_for('trash'))
    
    # Restaurar arquivo
    file.is_deleted = False
    file.deleted_at = None
    
    # Reativar compartilhamentos
    FileShare.query.filter_by(file_id=file_id).update({'is_active': True})
    
    db.session.commit()
    
    flash('Arquivo restaurado com sucesso!', 'success')
    return redirect(url_for('trash'))

@app.route('/permanent_delete/<int:file_id>', methods=['POST'])
@login_required
def permanent_delete_file(file_id):
    """Excluir arquivo permanentemente da lixeira"""
    file = File.query.get_or_404(file_id)
    
    # Verificar se o arquivo pertence ao usuário
    if file.user_id != current_user.id:
        flash('Você não tem permissão para excluir este arquivo!', 'error')
        return redirect(url_for('trash'))
    
    # Deletar todos os compartilhamentos relacionados ao arquivo
    FileShare.query.filter_by(file_id=file_id).delete()
    
    # Deletar arquivo físico se não for pasta
    if not file.is_folder and os.path.exists(file.file_path):
        try:
            os.remove(file.file_path)
        except Exception as e:
            print(f"Erro ao deletar arquivo físico: {e}")
    
    # Deletar do banco de dados
    db.session.delete(file)
    db.session.commit()
    
    flash('Arquivo excluído permanentemente!', 'success')
    return redirect(url_for('trash'))

@app.route('/empty_trash', methods=['POST'])
@login_required
def empty_trash():
    """Esvaziar toda a lixeira"""
    # Buscar todos os arquivos excluídos do usuário
    deleted_files = File.query.filter_by(
        user_id=current_user.id,
        is_deleted=True
    ).all()
    
    deleted_count = 0
    for file in deleted_files:
        try:
            # Deletar compartilhamentos
            FileShare.query.filter_by(file_id=file.id).delete()
            
            # Deletar arquivo físico se não for pasta
            if not file.is_folder and os.path.exists(file.file_path):
                os.remove(file.file_path)
            
            # Deletar do banco
            db.session.delete(file)
            deleted_count += 1
        except Exception as e:
            print(f"Erro ao deletar arquivo {file.id}: {e}")
    
    db.session.commit()
    
    flash(f'{deleted_count} arquivo(s) excluído(s) permanentemente!', 'success')
    return redirect(url_for('trash'))

@app.route('/api/toggle_favorite/<int:file_id>', methods=['POST'])
@login_required
def toggle_favorite(file_id):
    """Alternar status de favorito de um arquivo"""
    try:
        # Verificar permissão
        has_permission, permission, is_owner = check_file_permission(file_id, current_user.id, 'view')
        if not has_permission:
            return jsonify({'error': 'Sem permissão'}), 403
        
        file = File.query.get_or_404(file_id)
        file.is_favorite = not file.is_favorite
        db.session.commit()
        
        return jsonify({
            'success': True,
            'is_favorite': file.is_favorite
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/update_tags/<int:file_id>', methods=['POST'])
@login_required
def update_tags(file_id):
    """Atualizar tags de um arquivo"""
    try:
        # Verificar permissão
        has_permission, permission, is_owner = check_file_permission(file_id, current_user.id, 'edit')
        if not has_permission:
            return jsonify({'error': 'Sem permissão'}), 403
        
        tags = request.json.get('tags', '')
        file = File.query.get_or_404(file_id)
        file.tags = tags
        db.session.commit()
        
        return jsonify({
            'success': True,
            'tags': tags
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/move_file/<int:file_id>', methods=['POST'])
@login_required
def move_file(file_id):
    """Mover arquivo para outra pasta"""
    try:
        file = File.query.get_or_404(file_id)
        
        # Verificar se o usuário é o dono do arquivo
        if file.user_id != current_user.id:
            return jsonify({'error': 'Sem permissão para mover este arquivo'}), 403
        
        new_folder_id = request.json.get('folder_id')
        
        # Verificar se a pasta de destino existe e pertence ao usuário
        if new_folder_id:
            destination_folder = File.query.filter_by(
                id=new_folder_id, 
                user_id=current_user.id, 
                is_folder=True
            ).first()
            if not destination_folder:
                return jsonify({'error': 'Pasta de destino não encontrada'}), 404
        
        file.parent_folder_id = new_folder_id if new_folder_id else None
        db.session.commit()
        
        return jsonify({
            'success': True,
            'new_folder_id': new_folder_id
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/recent_files')
@login_required
def get_recent_files():
    """Obter arquivos acessados recentemente"""
    try:
        recent_files = File.query.filter_by(
            user_id=current_user.id,
            is_folder=False
        ).filter(
            File.last_accessed.isnot(None)
        ).order_by(File.last_accessed.desc()).limit(10).all()
        
        files_data = []
        for file in recent_files:
            files_data.append({
                'id': file.id,
                'filename': file.original_filename,
                'file_type': file.file_type,
                'file_size': format_file_size(file.file_size),
                'last_accessed': file.last_accessed.strftime('%d/%m/%Y %H:%M'),
                'is_favorite': file.is_favorite,
                'tags': file.tags or '',
                'url': url_for('view_file', file_id=file.id)
            })
        
        return jsonify({
            'files': files_data
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/favorite_files')
@login_required
def get_favorite_files():
    """Obter arquivos favoritos"""
    try:
        favorite_files = File.query.filter_by(
            user_id=current_user.id,
            is_folder=False,
            is_favorite=True
        ).order_by(File.updated_at.desc()).all()
        
        files_data = []
        for file in favorite_files:
            files_data.append({
                'id': file.id,
                'filename': file.original_filename,
                'file_type': file.file_type,
                'file_size': format_file_size(file.file_size),
                'is_favorite': file.is_favorite,
                'tags': file.tags or '',
                'url': url_for('view_file', file_id=file.id)
            })
        
        return jsonify({
            'files': files_data
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/user_folders')
@login_required
def get_user_folders():
    """Obter pastas do usuário para mover arquivos"""
    try:
        folders = File.query.filter_by(
            user_id=current_user.id,
            is_folder=True
        ).order_by(File.filename.asc()).all()
        
        folders_data = []
        for folder in folders:
            folders_data.append({
                'id': folder.id,
                'name': folder.original_filename
            })
        
        return jsonify({
            'folders': folders_data
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload_single', methods=['POST'])
@login_required
def upload_single_file():
    """API para upload de um único arquivo com resposta JSON"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'Nenhum arquivo selecionado'}), 400
        
        file = request.files['file']
        folder_id = request.form.get('folder_id', type=int)
        
        if file.filename == '':
            return jsonify({'success': False, 'error': 'Nome de arquivo vazio'}), 400
        
        # Criar pasta do usuário se não existir
        user_upload_folder = os.path.join(app.config['UPLOAD_FOLDER'], f'user_{current_user.id}')
        os.makedirs(user_upload_folder, exist_ok=True)
        
        # Verificar se a extensão é permitida
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'error': f'Extensão não permitida para {file.filename}'}), 400
        
        # Gerar nome único para o arquivo
        filename = safe_filename(file.filename)
        unique_filename = f"{uuid.uuid4()}_{filename}"
        file_path = os.path.join(user_upload_folder, unique_filename)
        
        # Salvar arquivo
        file.save(file_path)
        
        # Salvar no banco de dados
        db_file = File(
            filename=unique_filename,
            original_filename=filename,
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            file_type=os.path.splitext(filename)[1].lower(),
            user_id=current_user.id,
            parent_folder_id=folder_id
        )
        db.session.add(db_file)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'file_id': db_file.id,
            'filename': filename,
            'file_size': os.path.getsize(file_path)
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/download_folder/<int:folder_id>')
@login_required
def download_folder(folder_id):
    """Download de pasta compactada em ZIP"""
    try:
        # Verificar se a pasta existe e se o usuário tem permissão
        folder = File.query.filter_by(id=folder_id, is_folder=True).first()
        if not folder:
            flash('Pasta não encontrada.', 'error')
            return redirect(url_for('dashboard'))
        
        # Verificar permissão
        has_permission, permission, is_owner = check_file_permission(folder_id, current_user.id, 'view')
        if not has_permission:
            flash('Você não tem permissão para acessar esta pasta.', 'error')
            return redirect(url_for('dashboard'))
        
        # Criar diretório temporário para o ZIP
        import tempfile
        import zipfile
        
        temp_dir = tempfile.mkdtemp()
        zip_filename = f"{folder.original_filename}.zip"
        zip_path = os.path.join(temp_dir, zip_filename)
        
        # Criar arquivo ZIP
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            def add_folder_to_zip(folder_obj, base_path=""):
                """Adiciona recursivamente todos os arquivos da pasta ao ZIP"""
                # Buscar todos os arquivos e subpastas desta pasta
                children = File.query.filter_by(parent_folder_id=folder_obj.id, user_id=folder_obj.user_id).all()
                
                for child in children:
                    if child.is_folder:
                        # É uma subpasta, criar diretório e processar recursivamente
                        folder_path = os.path.join(base_path, child.original_filename)
                        add_folder_to_zip(child, folder_path)
                    else:
                        # É um arquivo, adicionar ao ZIP
                        if os.path.exists(child.file_path):
                            arcname = os.path.join(base_path, child.original_filename)
                            zipf.write(child.file_path, arcname)
            
            # Processar a pasta principal
            add_folder_to_zip(folder)
        
        # Enviar arquivo ZIP
        return send_file(
            zip_path,
            as_attachment=True,
            download_name=zip_filename,
            mimetype='application/zip'
        )
        
    except Exception as e:
        flash(f'Erro ao criar arquivo ZIP: {str(e)}', 'error')
        return redirect(url_for('dashboard'))
    finally:
        # Limpar arquivo temporário
        try:
            if 'zip_path' in locals():
                os.remove(zip_path)
            if 'temp_dir' in locals():
                shutil.rmtree(temp_dir)
        except:
            pass

# Variável global para armazenar progresso dos downloads
download_progress = {}

def download_youtube_thread(download_id, youtube_url, user_id, user_upload_folder, format_type='video', quality='best', is_playlist=False):
    """Função que roda em thread separada para fazer o download"""
    try:
        # Configurações do yt-dlp com callback de progresso
        def progress_hook(d):
            if d['status'] == 'downloading':
                if 'total_bytes' in d and d['total_bytes']:
                    progress = (d['downloaded_bytes'] / d['total_bytes']) * 100
                    download_progress[download_id]['progress'] = round(progress, 1)
                    download_progress[download_id]['message'] = f'Baixando... {download_progress[download_id]["progress"]}%'
                else:
                    download_progress[download_id]['message'] = 'Baixando...'
            elif d['status'] == 'finished':
                download_progress[download_id]['status'] = 'processing'
                download_progress[download_id]['message'] = 'Processando arquivo...'
                download_progress[download_id]['filename'] = d['filename']
                print(f"DEBUG: Arquivo finalizado: {d['filename']}")
            elif d['status'] == 'error':
                print(f"DEBUG: Erro no download: {d}")
                download_progress[download_id]['status'] = 'error'
                download_progress[download_id]['message'] = f'Erro no download: {d.get("error", "Erro desconhecido")}'
        
        # Configurar qualidade de vídeo (única opção disponível)
        if quality == 'best':
            format_spec = 'best[ext=mp4]/best'
        elif quality == '1080p':
            format_spec = 'best[height<=1080][ext=mp4]/best[height<=1080]/best'
        elif quality == '720p':
            format_spec = 'best[height<=720][ext=mp4]/best[height<=720]/best'
        elif quality == '480p':
            format_spec = 'best[height<=480][ext=mp4]/best[height<=480]/best'
        elif quality == '360p':
            format_spec = 'best[height<=360][ext=mp4]/best[height<=360]/best'
        else:
            format_spec = 'best[ext=mp4]/best'
        
        outtmpl = os.path.join(user_upload_folder, '%(title)s.%(ext)s')
        
        # Configurar ydl_opts
        ydl_opts = {
            'outtmpl': outtmpl,
            'format': format_spec,
            'noplaylist': not is_playlist,
            'progress_hooks': [progress_hook],
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            # Extrair informações do vídeo
            info = ydl.extract_info(youtube_url, download=False)
            video_title = info.get('title', 'video')
            download_progress[download_id]['title'] = video_title
            
            # Download do vídeo
            ydl.download([youtube_url])
            
            # Processar arquivo baixado
            if download_progress[download_id]['filename']:
                file_path = download_progress[download_id]['filename']
                filename = os.path.basename(file_path)
                
                # Debug: listar arquivos na pasta
                files_in_folder = os.listdir(user_upload_folder)
                print(f"DEBUG: Arquivos na pasta {user_upload_folder}: {files_in_folder}")
                print(f"DEBUG: Arquivo esperado: {filename}")
                print(f"DEBUG: Caminho completo: {file_path}")
                
                # Verificar se o arquivo existe
                if os.path.exists(file_path):
                    # Salvar no banco de dados (dentro do contexto da aplicação)
                    with app.app_context():
                        db_file = File(
                            filename=filename,
                            original_filename=filename,
                            file_path=file_path,
                            file_size=os.path.getsize(file_path),
                            file_type=os.path.splitext(filename)[1].lower(),
                            user_id=user_id
                        )
                        db.session.add(db_file)
                        db.session.commit()
                    
                    download_progress[download_id]['message'] = f'Vídeo "{video_title}" baixado com sucesso!'
                    
                    download_progress[download_id]['status'] = 'completed'
                    download_progress[download_id]['progress'] = 100
                else:
                    # Procurar por arquivos na pasta que correspondam ao título
                    found_file = None
                    for file_in_folder in files_in_folder:
                        file_in_folder_path = os.path.join(user_upload_folder, file_in_folder)
                        if os.path.isfile(file_in_folder_path):
                            # Verificar se o arquivo foi criado recentemente (últimos 30 segundos)
                            file_time = os.path.getctime(file_in_folder_path)
                            if datetime.now().timestamp() - file_time < 30:
                                # Verificar se é um arquivo de vídeo
                                if file_in_folder.lower().endswith(('.mp4', '.webm', '.mkv', '.avi', '.mov')):
                                    found_file = file_in_folder
                                    break
                    
                    if found_file:
                        file_path = os.path.join(user_upload_folder, found_file)
                        print(f"DEBUG: Arquivo encontrado: {found_file}")
                        
                        with app.app_context():
                            db_file = File(
                                filename=found_file,
                                original_filename=found_file,
                                file_path=file_path,
                                file_size=os.path.getsize(file_path),
                                file_type=os.path.splitext(found_file)[1].lower(),
                                user_id=user_id
                            )
                            db.session.add(db_file)
                            db.session.commit()
                        
                        download_progress[download_id]['message'] = f'Vídeo "{video_title}" baixado com sucesso!'
                        
                        download_progress[download_id]['status'] = 'completed'
                        download_progress[download_id]['progress'] = 100
                    else:
                        download_progress[download_id]['status'] = 'error'
                        download_progress[download_id]['message'] = f'Erro: Nenhum arquivo encontrado. Arquivos na pasta: {files_in_folder}'
            else:
                download_progress[download_id]['status'] = 'error'
                download_progress[download_id]['message'] = 'Erro ao processar arquivo baixado'
                
    except Exception as e:
        download_progress[download_id]['status'] = 'error'
        download_progress[download_id]['message'] = f'Erro: {str(e)}'

@app.route('/download_youtube', methods=['POST'])
@login_required
def download_youtube():
    """Download de vídeos do YouTube com progresso"""
    try:
        youtube_url = request.form.get('youtube_url')
        format_type = request.form.get('format', 'video')
        quality = request.form.get('quality', 'best')
        is_playlist = request.form.get('playlist') == 'on'
        
        if not youtube_url:
            return jsonify({'success': False, 'error': 'URL do YouTube é obrigatória!'})
        
        # Criar pasta do usuário se não existir
        user_upload_folder = os.path.join(app.config['UPLOAD_FOLDER'], f'user_{current_user.id}')
        os.makedirs(user_upload_folder, exist_ok=True)
        
        # Gerar ID único para este download
        download_id = str(uuid.uuid4())
        download_progress[download_id] = {
            'status': 'starting',
            'progress': 0,
            'message': 'Iniciando download...',
            'title': '',
            'filename': ''
        }
        
        # Iniciar download em thread separada
        thread = threading.Thread(
            target=download_youtube_thread,
            args=(download_id, youtube_url, current_user.id, user_upload_folder, format_type, quality, is_playlist)
        )
        thread.daemon = True
        thread.start()
        
        return jsonify({
            'success': True,
            'download_id': download_id,
            'message': 'Download iniciado com sucesso!'
        })
                
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Erro ao iniciar download: {str(e)}'
        })

@app.route('/youtube_progress/<download_id>')
@login_required
def youtube_progress(download_id):
    """Verificar progresso do download do YouTube"""
    if download_id in download_progress:
        return jsonify(download_progress[download_id])
    else:
        return jsonify({'status': 'not_found', 'message': 'Download não encontrado'})

@app.route('/youtube_info', methods=['POST'])
@login_required
def youtube_info():
    """Extrair informações do vídeo do YouTube"""
    try:
        youtube_url = request.form.get('youtube_url')
        if not youtube_url:
            return jsonify({'success': False, 'error': 'URL do YouTube é obrigatória!'})
        
        ydl_opts = {
            'quiet': True,
            'no_warnings': True,
            'extract_flat': False,
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(youtube_url, download=False)
            
            video_info = {
                'title': info.get('title', 'Título não disponível'),
                'duration': info.get('duration', 0),
                'uploader': info.get('uploader', 'Canal desconhecido'),
                'view_count': info.get('view_count', 0),
                'upload_date': info.get('upload_date', ''),
                'thumbnail': info.get('thumbnail', ''),
                'is_playlist': info.get('_type') == 'playlist',
                'playlist_count': info.get('playlist_count', 0) if info.get('_type') == 'playlist' else 0
            }
            
            return jsonify({
                'success': True,
                'info': video_info
            })
            
    except Exception as e:
        return jsonify({
            'success': False,
            'error': f'Erro ao extrair informações: {str(e)}'
        })

# ============================================================================
# ROTAS PARA ORGANIZAÇÕES E EQUIPES
# ============================================================================

@app.route('/organizations')
@login_required
def organizations():
    """Página principal de organizações"""
    # Buscar organizações que o usuário é dono ou membro
    owned_orgs = Organization.query.filter_by(owner_id=current_user.id).all()
    
    # Buscar organizações onde o usuário é membro de equipe
    member_orgs = db.session.query(Organization).join(Team).join(TeamMember).filter(
        TeamMember.user_id == current_user.id
    ).distinct().all()
    
    # Combinar e remover duplicatas
    all_orgs = list(set(owned_orgs + member_orgs))
    
    return render_template('organizations.html', organizations=all_orgs)

@app.route('/create_organization', methods=['GET', 'POST'])
@login_required
def create_organization():
    """Criar nova organização"""
    if request.method == 'POST':
        name = request.form.get('name')
        description = request.form.get('description')
        
        if not name:
            flash('Nome da organização é obrigatório!', 'error')
            return redirect(url_for('create_organization'))
        
        # Criar organização
        org = Organization(
            name=name,
            description=description,
            owner_id=current_user.id
        )
        db.session.add(org)
        db.session.commit()
        
        flash(f'Organização "{name}" criada com sucesso!', 'success')
        return redirect(url_for('organization_dashboard', org_id=org.id))
    
    return render_template('create_organization.html')

@app.route('/organization/<int:org_id>')
@login_required
def organization_dashboard(org_id):
    """Dashboard da organização"""
    org = Organization.query.get_or_404(org_id)
    
    # Verificar se o usuário tem acesso à organização
    if org.owner_id != current_user.id:
        # Verificar se é membro de alguma equipe
        is_member = db.session.query(TeamMember).join(Team).filter(
            Team.organization_id == org_id,
            TeamMember.user_id == current_user.id
        ).first()
        
        if not is_member:
            flash('Você não tem acesso a esta organização!', 'error')
            return redirect(url_for('organizations'))
    
    # Buscar equipes da organização
    teams = Team.query.filter_by(organization_id=org_id).all()
    
    # Buscar projetos da organização
    projects = Project.query.filter_by(organization_id=org_id).all()
    
    # Buscar pastas compartilhadas da organização
    shared_folders = SharedFolder.query.filter_by(organization_id=org_id).all()
    
    return render_template('organization_dashboard.html', 
                         organization=org, 
                         teams=teams, 
                         projects=projects, 
                         shared_folders=shared_folders)

@app.route('/organization/<int:org_id>/create_team', methods=['POST'])
@login_required
def create_team(org_id):
    """Criar nova equipe na organização"""
    org = Organization.query.get_or_404(org_id)
    
    # Verificar se o usuário é dono da organização
    if org.owner_id != current_user.id:
        flash('Apenas o dono da organização pode criar equipes!', 'error')
        return redirect(url_for('organization_dashboard', org_id=org_id))
    
    name = request.form.get('name')
    description = request.form.get('description')
    
    if not name:
        flash('Nome da equipe é obrigatório!', 'error')
        return redirect(url_for('organization_dashboard', org_id=org_id))
    
    # Criar equipe
    team = Team(
        name=name,
        description=description,
        organization_id=org_id
    )
    db.session.add(team)
    db.session.commit()
    
    flash(f'Equipe "{name}" criada com sucesso!', 'success')
    return redirect(url_for('organization_dashboard', org_id=org_id))

@app.route('/team/<int:team_id>/invite_user', methods=['POST'])
@login_required
def invite_user_to_team(team_id):
    """Convidar usuário para equipe"""
    team = Team.query.get_or_404(team_id)
    
    # Verificar se o usuário é dono da organização ou admin da equipe
    if team.organization.owner_id != current_user.id:
        # Verificar se é admin da equipe
        membership = TeamMember.query.filter_by(
            team_id=team_id,
            user_id=current_user.id,
            role='admin'
        ).first()
        
        if not membership:
            flash('Você não tem permissão para convidar usuários!', 'error')
            return redirect(url_for('organization_dashboard', org_id=team.organization_id))
    
    user_email = request.form.get('user_email')
    role = request.form.get('role', 'member')
    
    if not user_email:
        flash('Email do usuário é obrigatório!', 'error')
        return redirect(url_for('organization_dashboard', org_id=team.organization_id))
    
    # Buscar usuário pelo email
    user = User.query.filter_by(email=user_email).first()
    if not user:
        flash('Usuário não encontrado!', 'error')
        return redirect(url_for('organization_dashboard', org_id=team.organization_id))
    
    # Verificar se já é membro
    existing_member = TeamMember.query.filter_by(
        team_id=team_id,
        user_id=user.id
    ).first()
    
    if existing_member:
        flash('Usuário já é membro desta equipe!', 'error')
        return redirect(url_for('organization_dashboard', org_id=team.organization_id))
    
    # Adicionar membro
    member = TeamMember(
        team_id=team_id,
        user_id=user.id,
        role=role
    )
    db.session.add(member)
    db.session.commit()
    
    flash(f'Usuário {user.username} convidado com sucesso!', 'success')
    return redirect(url_for('organization_dashboard', org_id=team.organization_id))

@app.route('/organization/<int:org_id>/create_shared_folder', methods=['POST'])
@login_required
def create_shared_folder(org_id):
    """Criar pasta compartilhada na organização"""
    org = Organization.query.get_or_404(org_id)
    
    # Verificar se o usuário é dono da organização
    if org.owner_id != current_user.id:
        flash('Apenas o dono da organização pode criar pastas compartilhadas!', 'error')
        return redirect(url_for('organization_dashboard', org_id=org_id))
    
    name = request.form.get('name')
    description = request.form.get('description')
    
    if not name:
        flash('Nome da pasta é obrigatório!', 'error')
        return redirect(url_for('organization_dashboard', org_id=org_id))
    
    # Criar pasta compartilhada
    shared_folder = SharedFolder(
        name=name,
        description=description,
        organization_id=org_id,
        created_by=current_user.id
    )
    db.session.add(shared_folder)
    db.session.commit()
    
    flash(f'Pasta compartilhada "{name}" criada com sucesso!', 'success')
    return redirect(url_for('organization_dashboard', org_id=org_id))

@app.route('/shared_folder/<int:folder_id>')
@login_required
def shared_folder_view(folder_id):
    """Visualizar pasta compartilhada da organização"""
    shared_folder = SharedFolder.query.get_or_404(folder_id)
    
    # Verificar se o usuário tem acesso à organização
    if shared_folder.organization.owner_id != current_user.id:
        # Verificar se é membro de alguma equipe
        is_member = db.session.query(TeamMember).join(Team).filter(
            Team.organization_id == shared_folder.organization_id,
            TeamMember.user_id == current_user.id
        ).first()
        
        if not is_member:
            flash('Você não tem acesso a esta pasta!', 'error')
            return redirect(url_for('organizations'))
    
    # Buscar arquivos da pasta compartilhada
    files = File.query.filter_by(
        shared_folder_id=folder_id,
        is_deleted=False
    ).order_by(File.created_at.desc()).all()
    
    return render_template('shared_folder_view.html', 
                         shared_folder=shared_folder, 
                         files=files)

@app.route('/api/organization_users/<int:org_id>')
@login_required
def get_organization_users(org_id):
    """API para buscar usuários da organização"""
    org = Organization.query.get_or_404(org_id)
    
    # Verificar se o usuário tem acesso à organização
    if org.owner_id != current_user.id:
        is_member = db.session.query(TeamMember).join(Team).filter(
            Team.organization_id == org_id,
            TeamMember.user_id == current_user.id
        ).first()
        
        if not is_member:
            return jsonify({'error': 'Acesso negado'}), 403
    
    # Buscar todos os usuários que são membros de equipes da organização
    users = db.session.query(User).join(TeamMember).join(Team).filter(
        Team.organization_id == org_id
    ).distinct().all()
    
    # Adicionar o dono da organização se não estiver na lista
    if org.owner not in users:
        users.append(org.owner)
    
    users_data = [{'id': user.id, 'username': user.username, 'email': user.email} for user in users]
    
    return jsonify({'users': users_data})

@app.route('/organization/<int:org_id>/create_project', methods=['POST'])
@login_required
def create_project(org_id):
    """Criar novo projeto na organização"""
    org = Organization.query.get_or_404(org_id)
    
    # Verificar se o usuário é dono da organização
    if org.owner_id != current_user.id:
        flash('Apenas o dono da organização pode criar projetos!', 'error')
        return redirect(url_for('organization_dashboard', org_id=org_id))
    
    name = request.form.get('name')
    description = request.form.get('description')
    
    if not name:
        flash('Nome do projeto é obrigatório!', 'error')
        return redirect(url_for('organization_dashboard', org_id=org_id))
    
    # Criar projeto
    project = Project(
        name=name,
        description=description,
        organization_id=org_id,
        created_by=current_user.id
    )
    db.session.add(project)
    db.session.commit()
    
    flash(f'Projeto "{name}" criado com sucesso!', 'success')
    return redirect(url_for('organization_dashboard', org_id=org_id))

@app.route('/shared_folder/<int:folder_id>/upload', methods=['POST'])
@login_required
def upload_to_shared_folder(folder_id):
    """Upload de arquivo para pasta compartilhada"""
    shared_folder = SharedFolder.query.get_or_404(folder_id)
    
    # Verificar se o usuário tem acesso à organização
    if shared_folder.organization.owner_id != current_user.id:
        is_member = db.session.query(TeamMember).join(Team).filter(
            Team.organization_id == shared_folder.organization_id,
            TeamMember.user_id == current_user.id
        ).first()
        
        if not is_member:
            flash('Você não tem acesso a esta pasta!', 'error')
            return redirect(url_for('organizations'))
    
    if 'file' not in request.files:
        flash('Nenhum arquivo selecionado!', 'error')
        return redirect(url_for('shared_folder_view', folder_id=folder_id))
    
    file = request.files['file']
    if file.filename == '':
        flash('Nenhum arquivo selecionado!', 'error')
        return redirect(url_for('shared_folder_view', folder_id=folder_id))
    
    if file:
        # Verificar se a extensão é permitida
        if not allowed_file(file.filename):
            flash(f'Extensão não permitida para {file.filename}', 'error')
            return redirect(url_for('shared_folder_view', folder_id=folder_id))
        
        # Salvar arquivo
        filename = safe_filename(file.filename)
        unique_id = str(uuid.uuid4())
        new_filename = f"{unique_id}_{filename}"
        
        # Criar pasta da organização se não existir
        org_folder = os.path.join(app.config['UPLOAD_FOLDER'], f'org_{shared_folder.organization_id}')
        os.makedirs(org_folder, exist_ok=True)
        
        # Criar pasta da pasta compartilhada
        shared_folder_path = os.path.join(org_folder, f'folder_{folder_id}')
        os.makedirs(shared_folder_path, exist_ok=True)
        
        file_path = os.path.join(shared_folder_path, new_filename)
        file.save(file_path)
        
        # Salvar no banco de dados
        new_file = File(
            filename=new_filename,
            original_filename=filename,
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            file_type=os.path.splitext(filename)[1].lower(),
            user_id=current_user.id,
            shared_folder_id=folder_id
        )
        db.session.add(new_file)
        db.session.commit()
        
        flash(f'Arquivo "{filename}" enviado com sucesso!', 'success')
    
    return redirect(url_for('shared_folder_view', folder_id=folder_id))

@app.route('/shared_folder/<int:folder_id>/create_folder', methods=['POST'])
@login_required
def create_folder_in_shared(folder_id):
    """Criar pasta dentro da pasta compartilhada"""
    shared_folder = SharedFolder.query.get_or_404(folder_id)
    
    # Verificar se o usuário tem acesso à organização
    if shared_folder.organization.owner_id != current_user.id:
        is_member = db.session.query(TeamMember).join(Team).filter(
            Team.organization_id == shared_folder.organization_id,
            TeamMember.user_id == current_user.id
        ).first()
        
        if not is_member:
            flash('Você não tem acesso a esta pasta!', 'error')
            return redirect(url_for('organizations'))
    
    folder_name = request.form.get('folder_name')
    if not folder_name:
        flash('Nome da pasta é obrigatório!', 'error')
        return redirect(url_for('shared_folder_view', folder_id=folder_id))
    
    # Criar pasta no sistema de arquivos
    org_folder = os.path.join(app.config['UPLOAD_FOLDER'], f'org_{shared_folder.organization_id}')
    shared_folder_path = os.path.join(org_folder, f'folder_{folder_id}')
    new_folder_path = os.path.join(shared_folder_path, folder_name)
    
    if os.path.exists(new_folder_path):
        flash('Uma pasta com este nome já existe!', 'error')
        return redirect(url_for('shared_folder_view', folder_id=folder_id))
    
    os.makedirs(new_folder_path, exist_ok=True)
    
    # Salvar no banco de dados
    new_folder = File(
        filename=folder_name,
        original_filename=folder_name,
        file_path=new_folder_path,
        file_size=0,
        file_type='',
        is_folder=True,
        user_id=current_user.id,
        shared_folder_id=folder_id
    )
    db.session.add(new_folder)
    db.session.commit()
    
    flash(f'Pasta "{folder_name}" criada com sucesso!', 'success')
    return redirect(url_for('shared_folder_view', folder_id=folder_id))

def allowed_file(filename):
    """Verifica se a extensão do arquivo é permitida"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def safe_filename(filename):
    """Versão personalizada de secure_filename que preserva extensões importantes"""
    # Preservar a extensão original
    if '.' in filename:
        name, ext = filename.rsplit('.', 1)
        # Usar secure_filename apenas no nome, não na extensão
        safe_name = secure_filename(name)
        # Se o nome ficou vazio, usar um nome padrão
        if not safe_name:
            safe_name = 'file'
        return f"{safe_name}.{ext}"
    else:
        return secure_filename(filename)

if __name__ == '__main__':
    with app.app_context():
        # Criar tabelas se não existirem
        db.create_all()
        print("✅ Banco de dados inicializado!")
    
    app.run(debug=True, host='0.0.0.0', port=5000) 
